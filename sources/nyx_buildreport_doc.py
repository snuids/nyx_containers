import os
import re
import sys
import json
import pytz
import logging
import tzlocal
import datetime
import traceback
import cachetools
import collections
import numpy as np
import pandas as pd
from docx import Document
from datetime import timezone
from datetime import datetime
from docx.shared import Inches
import matplotlib.pyplot as plt
from dateutil.parser import parse
from docx.oxml import OxmlElement
from elastic_helper import es_helper
from docx.text.paragraph import Paragraph
#from docx.oxml.xmlchemy import OxmlElement
from docx.text.paragraph import Paragraph
from opensearchpy import OpenSearch as ES, RequestsHttpConnection as RC

logging.basicConfig()
logger = logging.getLogger()
logger.setLevel(logging.INFO)

#######################################################################################
# resize_table
#######################################################################################
def resize_table(table, size_array):
    table.autofit = False
    for i in range(len(size_array)):
        for cell in table.columns[i].cells:
            cell.width = Inches(size_array[i])

#######################################################################################
# save_log
#######################################################################################
def save_log():
    logger.info("Saving log")
    body={"@timestamp":datetime.now().isoformat(),"logs":logs}
    es.index("nyx_reportlog",id=report["id"],body=body)
    logger.info("Saved")

#######################################################################################
# insert_paragraph_after
#######################################################################################    
def insert_paragraph_after(inparagraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    inparagraph._p.addnext(new_p)
    
        
    new_para = Paragraph(new_p, inparagraph._parent)
    if text!=None:
        new_para.style=template.styles[style]
        run=new_para.add_run(text)
    return new_para

#######################################################################################
# compute style
#######################################################################################

def computestyle(level=1):
    style = "Title" if level == 0 else "Heading %d" % level    
    return style

#######################################################################################
# move_table_after
#######################################################################################

def move_table_after(table, paragraph):    
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

#######################################################################################
# create_table
#######################################################################################

def create_table( inpara,df, style=None,col_width=None):
    newpara=insert_paragraph_after(inpara,text="", style="Normal")
    global template
    table = template.add_table(df.shape[0]+1, df.shape[1])

    if style:
        table.style = style
        
    table_cells = table._cells

    for i in range(0,df.shape[0]+1):
        row_cells = table_cells[i*df.shape[1]:(i+1)*df.shape[1]]

        if i==0:
            for idx, col in enumerate(df.columns):
                row_cells[idx].text = col
        else:
            for idx, col in enumerate(df.columns):
                row_cells[idx].text = str(df.loc[(int(i-1)), col])


    inpara2=newpara.insert_paragraph_before()
    inpara2._p.addnext(table._tbl)
    
    if col_width!=None:
        table.autofit = False
        for i in range(len(col_width)):
            for cell in table.columns[i].cells:
                logger.info("resize %.2f" % col_width[i])
                cell.width = Inches(col_width[i])
    else:
        table.autofit = True

    return newpara,table


#######################################################################################
# replaceText
#######################################################################################
def replaceText(replacementHT,text):
    for keypair in replacementHT:
        try:
            text=text.replace("${"+keypair+"}",str(replacementHT[keypair])).replace('.0 %',' %')
        except:
            pass
    return text

################################################################################
def logger_info(log,exc_info=False):
    logger.info(log,exc_info=exc_info)
    logs.append("%s\tINFO\t%s" %(datetime.now().strftime("%d%b%Y %H:%M:%S.%f"),log))

def logger_error(log,exc_info=False):
    logger.error(log,exc_info=exc_info)
    logs.append("%s\tERROR\t%s" %(datetime.now().strftime("%d%b%Y %H:%M:%S.%f"),log))

def logger_debug(log,exc_info=False):
    logger.debug(log,exc_info=exc_info)
    logs.append("%s\tDEBUG\t%s" %(datetime.now().strftime("%d%b%Y %H:%M:%S.%f"),log))

def logger_warn(log,exc_info=False):
    logger.warn(log,exc_info=exc_info)
    logs.append("%s\tWARN\t%s" %(datetime.now().strftime("%d%b%Y %H:%M:%S.%f"),log))

def logger_fatal(log,exc_info=False):
    logger.fatal(log,exc_info=exc_info)
    logs.append("%s\tINFO\t%s" %(datetime.now().strftime("%d%b%Y %H:%M:%S.%f"),log))


#######################################################################################
# Start
#######################################################################################

print("Starting v0.2")
print(datetime.now(timezone.utc).isoformat())
localmode=False
try:
    os.environ["LOCAL_MODE"]
    localmode=True
except:
    pass

try:
    report=json.loads(sys.argv[1])
except:
    report={"id": "id_19485666_27963523", "creds": {"token": "933921e1-58ae-4948-a941-037e9ce2e915", "user": {"filters": [], "firstname": "Arnaud", "id": "amarchand@icloud.com", "language": "en", "lastname": "Marchand", "login": "amarchand", "password": "", "phone": "0033497441962", "privileges": ["admin"], "user": "amarchand@icloud.com"}}, "report": {"description": "My First Report", "exec": "report1", "generatePDF": True, "icon": "file", "output": ["txt", "pdf"], "parameters": [{"name": "param1", "title": "Param1", "type": "text", "value": "test"}, {"name": "param2", "title": "Param2", "type": "interval", "value": ["2019-11-20T23:00:00.000Z", "2019-11-21T23:00:00.000Z"]}], "privileges": [], "reportType": "python", "title": "My First Report"}, "privileges": ["admin"], "treatment": {"status": "Waiting", "creation": "2019-11-22T08:24:03.288Z", "start": "2019-11-22T09:24:03.293651"}, "@timestamp": "2019-11-22T08:24:03.288Z", "output": "/opt/sources/generated/id_19485666_27963523"}


logs=[]

functions="""def insert_paragraph_after(inparagraph, text=None, style=None)
def computestyle(level=1)
def create_table(paragraph,dataframe, style=None)
def resize_table(table, size_array)"""

logger_info("*"*90)
for f in functions.split("\n"):
    logger_info(f"* FUNCTION: {f}")
logger_info("*"*90)

params={}

containertimezone=tzlocal.get_localzone()
 
for param in report["report"]["parameters"]:
    if param["type"]=="interval":
        params[param["name"]+"_start"]=parse(param["value"][0]).astimezone(containertimezone)
        params[param["name"]+"_end"]=parse(param["value"][1]).astimezone(containertimezone)
    elif param["type"]=="date":
        params[param["name"]]=parse(param["value"]).astimezone(containertimezone)
    else:
        params[param["name"]]=param["value"]

if not localmode:
    prepath = "./reports/notebooks/"
else:
    prepath = "./reports/notebooks/"

logger_info("params=%s" %params)
logger_info("*"*90)
logger_info("LOCAL: <template> to reach the document")
logger_info("LOCAL: <paragraph> to reach a paragraph in a paragraph section")
logger_info("LOCAL: <params> to get parameters")
logger_info("LOCAL: <es> to reach elastic search")
logger_info("LOCAL: <replacementHT> to reach the replacement dictionary")

logger_info("*"*90)
logger_info("MARKUP: <#@ONLOAD> for load at startup cells")
logger_info("MARKUP: <#@PARAGRAPH=p_name> for paragraph cells")
logger_info("*"*90)
logger_info("VERSION: 1.0")
logger_info("*"*90)
# ELASTIC SEARCH

if not localmode:
    #host_params="http://10.0.10.161:9200"
    host_params="http://elasticsearch:9200"
    es = ES(hosts=[host_params])
    print(es.info())
else:        
    print("Connect to elastic via password")
    host_params = {'host':os.environ["ELK_URL"], 'port':int(os.environ["ELK_PORT"]), 'use_ssl':True}
    
    es = ES([host_params], connection_class=RC, http_auth=(os.environ["ELK_LOGIN"], os.environ["ELK_PASSWORD"]),  use_ssl=True ,verify_certs=False)
    print(es.info())

docpath=prepath+report["report"]["notebook"]+".docx"
ipynbpath=prepath+report["report"]["notebook"]+".ipynb"

# DOCS

if os.path.isfile(docpath) :
    print("DOC (%s) found." %(docpath))
else:
    print("ERROR DOC (%s) found." %(docpath))
    sys.exit(1) 

template = Document(docpath)

# PYTHON
if os.path.isfile(ipynbpath) :
    print("IPYNB (%s) found." %(ipynbpath))
else:
    print("ERROR IPYNB (%s) found." %(ipynbpath))
    sys.exit(1) 



replacementHT={}
replacementHT.update(params)




def computeNewCode(newcode):
    return newcode.replace("logger.","logger_")

reportfunctions={}

with open(ipynbpath, 'r') as content_file:
    content = content_file.read()
    jsoncontent=json.loads(content)
    for cell in jsoncontent["cells"]:
        if cell["cell_type"]=="code":
            if len(cell["source"])>0 and "#@ONLOAD" in cell["source"][0]:
                newcode="".join(cell["source"])
                try:
                    exec(computeNewCode(newcode))
                except:
                    logger_error("Common part crashed.",exc_info=True)
                    tb = traceback.format_exc()
                    for trace in ["TRACE:"+_ for _ in tb.split("\n")]:
                        logger_info(trace)
                    save_log()
                    raise Exception("ERROR COMMON")
            if len(cell["source"])>0 and "#@PARAGRAPH=" in cell["source"][0]:
                newcode="".join(cell["source"])
                reportfunctions["${"+cell["source"][0].replace("#@PARAGRAPH=","").strip()+"}"]=computeNewCode(newcode)
            

# FILL table cells
pattern = r"\${.*}"

for table in template.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:                            
                if re.search(pattern, paragraph.text):  
                    match= re.search(pattern, paragraph.text).group(0)
                    print(">>>>> FOUND CELL :%s" %(match))
                    logger_info(">>>> FOUND CELL :%s" %(match))  
                    if match in reportfunctions:
                        paragraph.text=paragraph.text.replace(match,"")
                        if match in reportfunctions:
                            logger_info(">> Code found")
                        try:
                            exec(reportfunctions[match])
                        except:
                            logger_error("Function crashed.",exc_info=True)
                            logger_info("====>"*30)
                            logger_info(reportfunctions[match])
                            logger_info("====>"*30)
                            tb = traceback.format_exc()
                            for trace in ["TRACE:"+_ for _ in tb.split("\n")]:
                                logger_info(trace)

                                if "<string>"  in trace and "<module>" in trace:
                                    try:
                                        linenbr=int(trace.split("line ")[1].split(",")[0])
                                        linecodes=newcode.split("\n")
                                        for i in range(linenbr-5,linenbr+4):
                                            if i>=0 and i<len(linecodes):
                                                if i==linenbr-1:
                                                    logger_error(">>> CODE:"+linecodes[i])
                                                else:
                                                    logger_info(">>> CODE:"+linecodes[i])

                                    except e:
                                        logger_error("ERROR while decoding error",exc_info=True)

                            save_log()
                            raise Exception("ERROR COMMON")
                    else:
                        paragraph.text=replaceText(replacementHT,paragraph.text)                   

# Fill paragraphs

for paragraph in template.paragraphs:     
    if re.search(pattern, paragraph.text):                
        match= re.search(pattern, paragraph.text).group(0)
        print(">>>>> FOUND PARAGRAPH :%s" %(match)) 
        logger_info(">>>>> FOUND PARAGRAPH :%s" %(match))                    
        paragraph.text=replaceText(replacementHT,paragraph.text)        
        if match in reportfunctions:
            paragraph.text=paragraph.text.replace(match,"")
            if match in reportfunctions:
                logger_info(">> Code found")

            try:
                exec(reportfunctions[match])
            except:
                logger_error("Function crashed.",exc_info=True)
                tb = traceback.format_exc()
                for trace in ["TRACE:"+_ for _ in tb.split("\n")]:
                    logger_info(trace)

                    if "<string>"  in trace and "<module>" in trace:
                           try:
                                linenbr=int(trace.split("line ")[1].split(",")[0])
                                linecodes=newcode.split("\n")
                                for i in range(linenbr-5,linenbr+4):
                                    if i>=0 and i<len(linecodes):
                                        if i==linenbr-1:
                                          logger_error(">>> CODE:"+linecodes[i])
                                        else:
                                          logger_info(">>> CODE:"+linecodes[i])

                           except e:
                                logger_error("ERROR while decoding error",exc_info=True)

                    
                save_log()
                raise Exception("ERROR COMMON")

# SAVE REPORT

if not localmode:
    template.save(report["output"]+".docx")
else:
    print('----save file in notebook mode----')
    template.save("notebook.docx")


save_log()
